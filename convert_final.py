import openpyxl
from docx import Document
import re

doc = Document('C:\\Users\\shave\\Documents\\Python soubory\\Programy\\Petra excel to word\\NEW\\BLANK.docx')

wb = openpyxl.load_workbook(filename = 
'C:\\Users\\shave\\Documents\\Python soubory\\Programy\\Petra excel to word\\NEW\\t1.xlsx')
b = wb.worksheets[0]
######################################################################
def ztabulky(row): #změna podle tabulky
    """
    row je číslo řádku
    """
    global nazev_soubor
    #měním Product name (T0,0,1) na "B{row}"
    a = doc.tables[0].cell(0,1)
    print(a.text)

    print(b[f"B{row}"].value)

    a.text = b[f"B{row}"].value
    print(a.text)

    nazev_soubor = b[f"B{row}"].value

    #měním Product Code (T0,1,1) na "A{row}"
    a = doc.tables[0].cell(1,1)
    print(a.text)

    print(b[f"A{row}"].value)
   
    a.text = b[f"A{row}"].value
    print(a.text)

    #změna Starting material - Quantity Required (T1,1,1 a T1,2,1) z "D"
    b1 = b[f"D{row}"].value
    print(b1)
    if b1==None:
        print("prázdné")
    else:
        c = re.findall("= ([\d\.g]+)", b1)
        #CBD Isolate 98%
        a = doc.tables[1].cell(2,1)
        a.text = c[0].strip()
        #Hempseed Oil
        a = doc.tables[1].cell(1,1)
        a.text = c[1].strip()

        #změna Starting material - Wastage (T1,1,2 a T1,2,2) na 7% z Quantity Required
        e = c[0].strip(" g")
        print(c)
        d = c[1].strip(" g")
        print(d)

        e1 = float(e)
        print(e1)
        d1 = float(d)
        print(d1)
        
        e2 = e1/100*7
        e2 = round(e2,1)
        print(e2)
        e2 = str(e2)
        a = doc.tables[1].cell(2,2)
        a.text = f"{e2}g"

        d2 = d1/100*7
        d2 = round(d2,1)
        print(d2)
        d2 = str(d2)
        a = doc.tables[1].cell(1,2)
        a.text = f"{d2}g"

    #T2,3,1 na % CBD z B
    a = doc.tables[2].cell(3,1)
    print(a.text)
    c = b[f"B{row}"].value
    d = re.findall("^[a-zA-Z\(\)\+ ]+([\d\.]+)%", c)
    print(d[0])
    a.text = f"{d[0]} %"


####################################################################
def opakovani(x,y):
    for i in range(x,y+1):#x=první řádek v tabulce z které čerpám text, y=poslední řádek
        ztabulky(i)
        doc.save(f'C:\\Users\\shave\\Documents\\Python soubory\\Programy\\Petra excel to word\\NEW\\t1_word\\{nazev_soubor}.docx')

####################################################################
opakovani(1001,1497) #použít na excel řádek (od, do)
